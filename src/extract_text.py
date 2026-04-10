# src/extract_text.py
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from pypdf import PdfReader
from docx import Document

from .normalize_med_text import normalize_med_text

SupportedType = Literal["pdf", "docx", "txt"]


@dataclass(frozen=True)
class ExtractedText:
    path: str
    file_type: SupportedType
    text: str


def extract_text(path: str, *, clinical: bool = False) -> ExtractedText:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Файл не найден: {path}")

    ext = p.suffix.lower().lstrip(".")
    if ext not in ("pdf", "docx", "txt"):
        raise ValueError(f"Неподдерживаемый формат: .{ext} (нужны pdf/docx/txt)")

    if ext == "pdf":
        text = _pdf_to_text(p)
        ftype: SupportedType = "pdf"
    elif ext == "docx":
        text = _docx_to_text(p)
        ftype = "docx"
    else:
        text = p.read_text(encoding="utf-8", errors="replace")
        ftype = "txt"

    # нормализация
    text, _ = normalize_med_text(text, clinical=clinical, unwrap_lines=True, return_log=False)

    # сигнал о возможном скане
    if ftype == "pdf" and len(text) < 200:
        text = (
            "[WARNING] Похоже, PDF — скан/картинки: извлечено слишком мало текста. "
            "Нужен OCR или текстовый PDF.\n\n" + text
        )

    return ExtractedText(path=str(p), file_type=ftype, text=text)


def _pdf_to_text(p: Path) -> str:
    reader = PdfReader(str(p))
    chunks: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        t = (page.extract_text() or "").strip()
        if t:
            chunks.append(f"[PAGE {i}]\n{t}")
    return "\n\n".join(chunks)


def _docx_to_text(p: Path) -> str:
    doc = Document(str(p))
    parts: list[str] = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    # таблицы (очень важно для ИГХ/патологии)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = (cell.text or "").strip().replace("\n", " ")
                if ct:
                    cells.append(ct)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
